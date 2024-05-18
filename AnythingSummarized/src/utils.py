from PyPDF2 import PdfReader
import docx
import requests
from ebooklib import epub
import numpy as np
from os import environ as osenv
from docx import Document
from glob import glob
import re
import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi
from bs4 import BeautifulSoup
from newspaper import Article
import base64
import whisper
WHISPER_MODEL = 'base'

def extract_youtube_video_id(url):
    patterns = [
        r"v=([a-zA-Z0-9_-]{11})",
        r"youtu\.be/([a-zA-Z0-9_-]{11})",
        r"youtube\.com/shorts/([a-zA-Z0-9_-]{11})",
        r"youtube\.com/embed/([a-zA-Z0-9_-]{11})"
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    raise ValueError("Invalid YouTube URL")

# Define the function to display file preview in the Streamlit app
def display_file_preview(uploaded_file):
    if uploaded_file.type == "application/pdf":
        pdf_file = uploaded_file.getvalue()
        base64_pdf = base64.b64encode(pdf_file).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="400" style="border:none;"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        full_text = '\n'.join([para.text for para in doc.paragraphs if para.text])
        st.text_area("DOCX Content", full_text, height=300)
    elif uploaded_file.type.startswith("text/"):
        text_string = str(uploaded_file.read(), "utf-8")
        st.text_area("Text Content", text_string, height=300)

def display_link_preview(link):
    # Check if the link is a YouTube video or an article
    if "youtube.com" in link or "youtu.be" in link:
        # Display YouTube video
        video_id = extract_youtube_video_id(link)
        video_url = f"https://www.youtube.com/watch?v={video_id}"
        st.video(video_url)

    
    elif link.endswith('.pdf'):
        pdf_display = f'<iframe src="{link}" width="100%" height="400" style="border:none;"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)
    else:
        # Display article content
        response = requests.get(link)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            article_text = soup.get_text()
            st.text_area("Article Content", article_text, height=300)

        else:
            st.error("Failed to fetch the article")
            return ""

# Function to extract text based on content type
def extract_text(input_data):
    if isinstance(input_data, str) and ("youtube.com" in input_data or "youtu.be" in input_data):
        video_id = extract_youtube_video_id(input_data)
        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
        return ' '.join([t['text'] for t in transcript])
    elif isinstance(input_data, str) and input_data.startswith('http'):
        article = Article(input_data)
        article.download()
        article.parse()
        return article.text
    else:
        # # Handle UploadedFile object directly
        file_extension = input_data.name.split('.')[-1].lower()
        print(file_extension)
        # if file_extension == "mp4":
        #     vid = input_data.name
        #     with open(vid, mode="wb") as f:
        #         f.write(input_data.read())
        #     base_dir = st.text_input('Set your videos base path', osenv.get('HOMEPATH'))
        #     path = os.path.join(base_dir)
        #     files = glob(path + r"\*.mp4") + glob(path + r"\*.avi")
         
        #     video_tensor = torchvision.io.read_video(files[2])
        #                 # Assuming video_tensor is a tuple and you need to extract the relevant array
        #     video_array = video_tensor[0]  # Assuming the array is at the first position in the tuple

        #     # Convert the array to a NumPy array
        #     video_array_np = np.array(video_array)

        #     # Now, pass the NumPy array to the transcribe function
           
        #     print(video_array_np)
        #     model = whisper.load_model(WHISPER_MODEL)
        #     transcript = model.transcribe(video_array_np)
        #     transcript = transcript['text']
            
        #     return transcript
        # Determine file type based on extension
        if file_extension == "pdf":
            # Extract text from PDF file
            Pdf_reader = PdfReader(input_data)
            text = ""
            for page in Pdf_reader.pages:
                text += page.extract_text()
            return text
        elif file_extension == "docx":
            # Extract text from DOCX file
            doc = Document(input_data)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        elif file_extension == "txt":
            # Extract text from TXT file
            text = input_data.getvalue().decode('utf-8')
            return text
        elif file_extension == "csv":
            # Extract text from CSV file
            csv_data = input_data.getvalue().decode('utf-8')
    
            return csv_data
        elif file_extension == "epub":
            # Extract text from EPUB file
            book = epub.read_epub(input_data)
            text = ""
            for item in book.get_items():
                if item.get_type() == epub.EpubHtml:
                    text += item.get_content().decode("utf-8")
            return text
        else:
            return "Unsupported file type."